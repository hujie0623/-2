
import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
import datetime
import os

st.set_page_config(page_title="旅游报价生成器", layout="centered")
st.title("📄 旅游报价一键生成器（增强版）")
st.markdown("填写行程信息，上传LOGO与每日行程段落，点击生成完整文档")

with st.form("quote_form"):
    col1, col2 = st.columns(2)
    with col1:
        title = st.text_input("行程标题", "北欧三国12天游")
        city = st.text_input("出发地", "深圳（香港起飞）")
        date = st.text_input("出发日期", "2025年8月15日")
        days = st.number_input("行程天数", 1, 30, 12)
        pax = st.number_input("人数", 1, 100, 10)
    with col2:
        dest = st.text_input("目的地（顿号分隔）", "丹麦、挪威、芬兰")
        hotel = st.text_input("酒店标准", "四星/部分升级五星")
        feature = st.text_input("特色体验（顿号分隔）", "峡湾游船、极光玻璃屋、圣诞老人村")
        price = st.text_input("预计报价", "约17100元/人")
        logo_file = st.file_uploader("上传LOGO图片 (PNG/JPG)", type=["png", "jpg"])

    st.markdown("---")
    st.subheader("📑 上传每日行程段落（可选 .docx）")
    daily_doc = st.file_uploader("上传包含每日行程段落的Word文档（可选）", type=["docx"])

    st.markdown("---")
    st.subheader("📷 上传每日插图（最多3张，将横向排列）")
    images = st.file_uploader("上传图片（可选，推荐3张）", type=["jpg", "png"], accept_multiple_files=True)

    submitted = st.form_submit_button("✅ 生成完整报价方案文档")

def init_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(11)
    return doc

def add_logo(doc, logo_stream):
    section = doc.sections[0]
    header = section.header
    if logo_stream:
        header.paragraphs[0].add_run().add_picture(logo_stream, width=Inches(1))

def add_cover(doc):
    doc.add_heading(f"梦幻北欧 · {title}", level=1)
    doc.add_paragraph(f"{dest}｜峡湾极光｜纯净童话｜全程纯玩")
    doc.add_heading("【产品亮点】", level=2)
    doc.add_paragraph(
        f"✓ 一次游览：{dest}\n"
        f"✓ 特色项目：{feature}\n"
        f"✓ 酒店安排：{hotel}\n"
        f"✓ 全程无购物无自费\n"
        f"✓ 出发时间：{date}｜{pax}人小团"
    )

def add_daily_details(doc, daily_doc, images):
    doc.add_heading("【每日行程详解】", level=2)
    if daily_doc:
        daily_docx = Document(daily_doc)
        for para in daily_docx.paragraphs:
            doc.add_paragraph(para.text)
            doc.add_paragraph("")
    else:
        for i in range(1, 4):
            doc.add_heading(f"DAY {i}｜示例日程", level=3)
            doc.add_paragraph("🚌 行程安排：参考景点游览 + 城市穿越 + 用车说明")
            doc.add_paragraph("📍 景点亮点：")
            doc.add_paragraph("• 核心景点一
• 景点二
• 景点三")
            doc.add_paragraph("🍽 餐食与住宿：中餐 + 晚餐 + 四星酒店")
    if images:
        p = doc.add_paragraph("📷 插图建议：")
        for img in images:
            run = p.add_run()
            run.add_picture(img, width=Inches(1.7))

def add_price_table(doc):
    doc.add_heading("【报价明细】", level=2)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Light Grid'
    headers = ['费用项目', '单价（元/人）', '说明']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
    data = [
        ['国际机票', '6800', '含税多段'],
        ['住宿费用', '5800', '四星起步'],
        ['交通导服', '1800', '含用车/导游/服务费'],
        ['门票餐食', '2700', '全含套餐']
    ]
    for row in data:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val

def add_footer(doc):
    doc.add_heading("【费用包含】", level=2)
    doc.add_paragraph("机票、酒店、导游、交通、门票、签证、保险等")
    doc.add_heading("【费用不含】", level=2)
    doc.add_paragraph("护照办理、单房差、个人消费")
    doc.add_heading("【报价信息】", level=2)
    doc.add_paragraph(f"报价：{price}\n适用人数：{pax}人\n出发时间：{date}\n报价时间：{datetime.datetime.now().strftime('%Y年%m月')}\n联系人：小游\n电话：138-0000-0000")

if submitted:
    doc = init_doc()
    add_logo(doc, logo_file)
    add_cover(doc)
    add_daily_details(doc, daily_doc, images)
    add_price_table(doc)
    add_footer(doc)
    filename = f"{title.replace(' ', '_')}_报价方案.docx"
    doc.save(filename)
    with open(filename, "rb") as f:
        st.success("✅ 报价方案生成成功！")
        st.download_button("📥 点击下载 Word 文件", f, file_name=filename, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    os.remove(filename)
